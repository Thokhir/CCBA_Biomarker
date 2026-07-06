"""Renders a ReportContent object to Word (.docx) via python-docx."""
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parents[2]
RESULTS_DIR = BASE_DIR / "results"
TEAL = RGBColor(0x0F, 0x6E, 0x68)


def _add_dataframe_table(doc, df) -> None:
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Light Grid Accent 1"
    header_cells = table.rows[0].cells
    for i, column in enumerate(df.columns):
        header_cells[i].text = str(column)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = f"{value:.4f}" if isinstance(value, float) else str(value)


def render_docx(content, out_path: Path) -> None:
    doc = Document()

    title = doc.add_heading(content.title, level=0)
    title.runs[0].font.color.rgb = TEAL

    subtitle = doc.add_paragraph(f"{content.subtitle} · Generated {content.generated_at}")
    subtitle.runs[0].italic = True

    for section in content.sections:
        heading = doc.add_heading(section.title, level=1)
        heading.runs[0].font.color.rgb = TEAL

        for paragraph in section.paragraphs:
            doc.add_paragraph(paragraph)

        if section.table is not None:
            _add_dataframe_table(doc, section.table)
            if section.table_caption:
                caption = doc.add_paragraph(section.table_caption)
                caption.runs[0].italic = True

        if section.figure_path:
            fig_file = RESULTS_DIR / section.figure_path
            if fig_file.exists():
                doc.add_picture(str(fig_file), width=Inches(5.5))
                if section.figure_caption:
                    caption = doc.add_paragraph(section.figure_caption)
                    caption.runs[0].italic = True
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(out_path))
